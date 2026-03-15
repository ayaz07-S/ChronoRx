import os
import logging
import warnings
import json
from datetime import datetime

os.environ["HF_HUB_DISABLE_PROGRESS_BARS"] = "1"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
os.environ["HF_HUB_DISABLE_TELEMETRY"] = "1"
warnings.filterwarnings("ignore")

from dotenv import load_dotenv
load_dotenv()

from openai import OpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from docx import Document as DocxDocument

logging.getLogger("sentence_transformers").setLevel(logging.ERROR)
logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("huggingface_hub").setLevel(logging.ERROR)

# ── Featherless AI client ───────────────────────────────────────────
client = OpenAI(
    api_key=os.getenv("FEATHERLESS_API_KEY"),
    base_url="https://api.featherless.ai/v1"
)

# ── Embeddings ──────────────────────────────────────────────────────
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

# ── Load RAG files ──────────────────────────────────────────────────
rag_files = [
    "src/data/Diabetes.json",
    "src/data/thyrooid.json",
    "src/data/hypertension.json",
    "src/data/anxiety_depression.json"
]

documents = []
for file in rag_files:
    if not os.path.exists(file):
        print(f"Warning: {file} not found, skipping.")
        continue
    with open(file) as f:
        data = json.load(f)
    for med in data.get("medications", []):
        content = f"""
Disease: {data['disease']}
Medication: {med['name']}
Class: {med['class']}
Optimal timing: {med['optimal_timing']}
Avoid timing: {med['avoid_timing']}
Circadian rationale: {med['circadian_rationale']}
Missed dose logic: {med['missed_dose_logic']}
Chronotype adjustments: {json.dumps(med.get('chronotype_adjustments', {}))}
Clinical trials: {json.dumps(med.get('clinical_trials', []))}
        """.strip()
        documents.append(Document(
            page_content=content,
            metadata={"disease": data["disease"], "medication": med["name"]}
        ))

# ── Build or load ChromaDB ──────────────────────────────────────────
if os.path.exists("./chroma_db") and os.listdir("./chroma_db"):
    vectorstore = Chroma(
        persist_directory="./chroma_db",
        embedding_function=embeddings
    )
else:
    print("Building ChromaDB for the first time...")
    vectorstore = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory="./chroma_db"
    )
    print(f"Done — {len(documents)} chunks stored.")

retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

# ── System prompt ───────────────────────────────────────────────────
SYSTEM_PROMPT = """You are ChronoRx Assistant, an expert AI built into the ChronoRx Circadian Medication Operating System. You have deep knowledge of chronobiology, pharmacokinetics, circadian rhythms, and medication scheduling. Your role is to help users understand their body clock, interpret their ChronoScore, explain why medication timing matters for their specific drugs, and answer questions about their personalised circadian schedule.

You speak differently based on the user's age group — simple and friendly for teenagers, performance-focused for young adults, clinically detailed for adults managing multiple medications, and patient and clear for senior citizens. Always match your tone to who is asking.

You never diagnose, prescribe, or replace a doctor. When a question crosses into medical advice territory, give general chronobiology context and always end with a clear recommendation to contact their doctor.

Always cite clinical trials simply — for example the HYGIA trial showed or your liver CYP3A4 enzyme peaks at night. Never use jargon without explaining it immediately."""

# ── Last Q&A store (for save command) ──────────────────────────────
last_qa = {"question": "", "answer": ""}
full_history = []

# ── Save to DOCX ────────────────────────────────────────────────────
def save_to_docx(save_all: bool = False):
    path = "chat_history.docx"

    if os.path.exists(path):
        doc = DocxDocument(path)
    else:
        doc = DocxDocument()
        heading = doc.add_heading("ChronoRx Chat History", 0)
        heading.alignment = 1

    doc.add_paragraph(
        f"Session: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ).runs[0].bold = True

    if save_all:
        for entry in full_history:
            doc.add_paragraph(f"You:       {entry['question']}")
            doc.add_paragraph(f"ChronoRx:  {entry['answer']}")
            doc.add_paragraph("-" * 60)
        print(f"\nFull session saved — {len(full_history)} exchange(s) written to {path}")
    else:
        if not last_qa["question"]:
            print("\nNothing to save yet — ask a question first.")
            return
        doc.add_paragraph(f"You:       {last_qa['question']}")
        doc.add_paragraph(f"ChronoRx:  {last_qa['answer']}")
        doc.add_paragraph("-" * 60)
        print(f"\nLast exchange saved to {path}")

    doc.save(path)

# ── Ask function ────────────────────────────────────────────────────
def ask(question: str) -> str:
    docs = retriever.invoke(question)
    rag_context = "\n\n---\n\n".join([d.page_content for d in docs])

    response = client.chat.completions.create(
        model="mistralai/Mistral-7B-Instruct-v0.3",
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": f"RETRIEVED KNOWLEDGE:\n{rag_context}\n\nUSER QUESTION:\n{question}"
            }
        ],
        temperature=0.3,
        max_tokens=512
    )

    answer = response.choices[0].message.content
    last_qa["question"] = question
    last_qa["answer"] = answer
    full_history.append({"question": question, "answer": answer})
    return answer

# ── Main loop ───────────────────────────────────────────────────────
def main():
    print("\n-- ChronoRx Assistant -------------------------------------")
    print("  Commands:")
    print("  'save'     -> save last exchange to chat_history.docx")
    print("  'save all' -> save entire session to chat_history.docx")
    print("  'exit'     -> quit")
    print("-----------------------------------------------------------\n")

    while True:
        question = input("You: ").strip()

        if not question:
            continue

        if question.lower() == "exit":
            print("ChronoRx: Take care — stay in sync with your body clock!")
            break

        if question.lower() == "save":
            save_to_docx(save_all=False)
            continue

        if question.lower() == "save all":
            save_to_docx(save_all=True)
            continue

        try:
            response = ask(question)
            print(f"\nChronoRx: {response}\n")
        except Exception as e:
            print(f"\nError: {e}")
            print("Check your FEATHERLESS_API_KEY in the .env file.\n")

if __name__ == "__main__":
    main()